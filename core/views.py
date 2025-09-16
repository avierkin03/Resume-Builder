from django.views.generic import TemplateView, FormView, CreateView, UpdateView, DeleteView, ListView, DetailView
from django.contrib.auth.views import LoginView, LogoutView
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib.auth.models import User, Group
from django.urls import reverse_lazy
from django.shortcuts import get_object_or_404, redirect
from django.http import HttpResponse, JsonResponse
from django.contrib import messages
from django.db.models import Q
from .forms import RegistrationForm, ResumeForm, SectionFormSet, get_section_formset
from .models import Resume, ResumeTemplate, ResumeSection, Announcement, Profile
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph
from reportlab.lib.styles import ParagraphStyle
from docx import Document
from docx.shared import Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import io
from django.contrib.auth import login
from django.template import engines
import os
from django.conf import settings
from django.core.paginator import Paginator
from rest_framework import generics
from .serializers import *


# В'ю для домашньої сторінки
class HomeView(TemplateView):
    template_name = 'home.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['templates'] = ResumeTemplate.objects.all()[:2]  # Популярні шаблони
        context['announcements'] = Announcement.objects.filter(is_active=True).order_by('-created_at')[:5]
        return context


# AJAX-view для пагінації шаблонів
def template_ajax(request):
    templates = ResumeTemplate.objects.all().order_by('-created_at')  # Сортуємо за датою створення
    page_number = request.GET.get('page', 1)
    paginator = Paginator(templates, 2)  # 3 шаблони на сторінку
    page_obj = paginator.get_page(page_number)
    data = [{
        'id': template.id,
        'name': template.name,
        'description': template.description,
        'image': template.image.url if template.image else None
    } for template in page_obj]
    return JsonResponse({
        'templates': data,
        'has_previous': page_obj.has_previous(),
        'has_next': page_obj.has_next(),
        'previous_page': page_obj.previous_page_number() if page_obj.has_previous() else None,
        'next_page': page_obj.next_page_number() if page_obj.has_next() else None,
        'num_pages': paginator.num_pages,
    })


# В'ю для реєстрації нового користувача
class RegisterView(FormView):
    template_name = 'register.html'
    form_class = RegistrationForm
    success_url = reverse_lazy('home')

    def form_valid(self, form):
        user = form.save()
        group = Group.objects.get_or_create(name='users')[0]
        user.groups.add(group)
        login(self.request, user)
        messages.success(self.request, "Реєстрація успішна!")
        return super().form_valid(form)


# В'ю для логіну користувача
class CustomLoginView(LoginView):
    template_name = 'login.html'
    redirect_authenticated_user = True

    def get_success_url(self):
        messages.success(self.request, "Вхід успішний!")
        return reverse_lazy('home')


# В'ю для оновлення профілю користувача
class ProfileView(LoginRequiredMixin, UpdateView):
    model = Profile
    fields = ['bio']
    template_name = 'profile.html'
    success_url = reverse_lazy('profile')

    def get_object(self):
        profile, created = Profile.objects.get_or_create(user=self.request.user)
        return profile

    def form_valid(self, form):
        messages.success(self.request, "Профіль оновлено!")
        return super().form_valid(form)


# В'ю для списку всіх резюме користувача
class ResumeListView(LoginRequiredMixin, ListView):
    model = Resume
    template_name = 'resume_list.html'
    context_object_name = 'resumes'

    def get_queryset(self):
        return Resume.objects.filter(user=self.request.user).order_by('-updated_at')


# В'ю для створення нового резюме користувача
class ResumeCreateView(LoginRequiredMixin, CreateView):
    model = Resume
    form_class = ResumeForm
    template_name = 'resume_form.html'
    success_url = reverse_lazy('resumes')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['section_formset'] = get_section_formset(
            resume=None,
            data=self.request.POST or None,
            files=self.request.FILES or None
        )
        template_id = self.request.GET.get('template_id')
        if template_id:
            context['form'].initial['template'] = template_id
        return context

    def form_valid(self, form):
        form.instance.user = self.request.user
        self.object = form.save()
        section_formset = get_section_formset(
            resume=self.object,
            data=self.request.POST,
            files=self.request.FILES
        )
        if section_formset.is_valid():
            section_formset.save()
            messages.success(self.request, "Резюме створено!")
            return super().form_valid(form)
        else:
            messages.error(self.request, f"Помилка при збереженні секцій: {section_formset.errors}")
            return self.form_invalid(form)

    def form_invalid(self, form):
        print("Form invalid. ResumeForm errors:", form.errors)
        print("SectionFormSet errors:", self.request.POST.get('section_formset', 'No section_formset data'))
        messages.error(self.request, f"Помилка при створенні резюме: {form.errors}")
        return super().form_invalid(form)


# В'ю для оновлення існуючого резюме користувача
class ResumeUpdateView(LoginRequiredMixin, UpdateView):
    model = Resume
    form_class = ResumeForm
    template_name = 'resume_form.html'
    success_url = reverse_lazy('resumes')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['section_formset'] = get_section_formset(
            resume=self.object,
            data=self.request.POST or None,
            files=self.request.FILES or None
        )
        return context

    def form_valid(self, form):
        self.object = form.save()
        section_formset = get_section_formset(
            resume=self.object,
            data=self.request.POST,
            files=self.request.FILES
        )
        if section_formset.is_valid():
            # Перевірка унікальності порядку
            orders = [f.cleaned_data.get('order') for f in section_formset if f.cleaned_data.get('order') is not None]
            if len(orders) != len(set(orders)):
                messages.error(self.request, "Порядок секцій має бути унікальним.")
                return self.form_invalid(form)

            section_formset.save()
            messages.success(self.request, "Резюме оновлено!")
            return super().form_valid(form)
        else:
            messages.error(self.request, f"Помилка при збереженні секцій: {section_formset.errors}")
            return self.form_invalid(form)


# В'ю для видалення існуючого резюме користувача
class ResumeDeleteView(LoginRequiredMixin, DeleteView):
    model = Resume
    template_name = 'resume_confirm_delete.html'
    success_url = reverse_lazy('resumes')

    def get_queryset(self):
        return Resume.objects.filter(user=self.request.user)

    def form_valid(self, form):
        messages.success(self.request, "Резюме видалено!")
        return super().form_valid(form)


# В'ю для клонування існуючого резюме користувача
class ResumeCloneView(LoginRequiredMixin, DetailView):
    model = Resume
    template_name = 'resume_list.html'  # Не потрібна окрема сторінка

    def get(self, request, *args, **kwargs):
        resume = self.get_object()
        if resume.user != self.request.user:
            messages.error(self.request, "Доступ заборонено!")
            return redirect('resumes')
        new_resume = Resume.objects.create(
            user=resume.user,
            title=f"Copy of {resume.title}",
            template=resume.template,
            photo=resume.photo
        )
        for section in resume.sections.all():
            ResumeSection.objects.create(
                resume=new_resume,
                section_type=section.section_type,
                content=section.content,
                order=section.order
            )
        messages.success(self.request, "Резюме склоновано!")
        return redirect('resumes')


# В'ю для прев'ю конкретного резюме
class ResumePreviewView(LoginRequiredMixin, DetailView):
    model = Resume
    template_name = 'resume_preview.html'
    context_object_name = 'resume'

    def get_queryset(self):
        return Resume.objects.filter(user=self.request.user)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        if self.object.template:
            template_engine = engines['django']
            template = template_engine.from_string(self.object.template.html_template)
            context['rendered_template'] = template.render({
                'resume': self.object,
                'sections': self.object.sections.all(),
            })
        return context


# В'ю для списку шаблонів резюме
class TemplateListView(ListView):
    model = ResumeTemplate
    template_name = 'template_list.html'
    context_object_name = 'templates'


# В'ю для списку оголошень
class AnnouncementListView(ListView):
    model = Announcement
    template_name = 'announcement_list.html'
    context_object_name = 'announcements'

    def get_queryset(self):
        return Announcement.objects.filter(is_active=True).order_by('-created_at')


# В'ю для створення нового оголошення
class AnnouncementCreateView(LoginRequiredMixin, CreateView):
    model = Announcement
    fields = ['title', 'content', 'is_active']
    template_name = 'announcement_form.html'
    success_url = reverse_lazy('announcements')

    def get_queryset(self):
        return Announcement.objects.filter(created_by=self.request.user)

    def form_valid(self, form):
        if not self.request.user.groups.filter(name='admins').exists():
            messages.error(self.request, "Тільки адміністратори можуть створювати оголошення!")
            return redirect('announcements')
        form.instance.created_by = self.request.user
        messages.success(self.request, "Оголошення створено!")
        return super().form_valid(form)


# В'ю для редагування існуючого оголошення
class AnnouncementUpdateView(LoginRequiredMixin, UpdateView):
    model = Announcement
    fields = ['title', 'content', 'is_active']
    template_name = 'announcement_form.html'
    success_url = reverse_lazy('announcements')

    def get_queryset(self):
        return Announcement.objects.filter(created_by=self.request.user)

    def form_valid(self, form):
        if not self.request.user.groups.filter(name='admins').exists():
            messages.error(self.request, "Тільки адміністратори можуть редагувати оголошення!")
            return redirect('announcements')
        messages.success(self.request, "Оголошення оновлено!")
        return super().form_valid(form)


# В'ю для видалення існуючого оголошення
class AnnouncementDeleteView(LoginRequiredMixin, DeleteView):
    model = Announcement
    template_name = 'announcement_confirm_delete.html'
    success_url = reverse_lazy('announcements')

    def get_queryset(self):
        return Announcement.objects.filter(created_by=self.request.user)

    def form_valid(self, form):
        if not self.request.user.groups.filter(name='admins').exists():
            messages.error(self.request, "Тільки адміністратори можуть видаляти оголошення!")
            return redirect('announcements')
        messages.success(self.request, "Оголошення видалено!")
        return super().form_valid(form)


# В'ю для експорту конкретного резюме в pdf-форматі
def export_pdf(request, pk):
    resume = get_object_or_404(Resume, pk=pk, user=request.user)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{resume.title}.pdf"'
    
    # Створюємо PDF
    p = canvas.Canvas(response, pagesize=A4)
    width, height = A4  # Розміри сторінки: 595.27 x 841.89 pt
    
    # Налаштування шрифтів
    font_path = os.path.join(settings.BASE_DIR, 'static', 'fonts', 'DejaVuSans.ttf')
    font_bold_path = os.path.join(settings.BASE_DIR, 'static', 'fonts', 'DejaVuSans-Bold.ttf')
    try:
        pdfmetrics.registerFont(TTFont('DejaVuSans', font_path))
        pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', font_bold_path))
        font_name = 'DejaVuSans'
        font_bold_name = 'DejaVuSans-Bold'
    except Exception as e:
        print(f"Error loading font: {e}")
        font_name = 'Helvetica'
        font_bold_name = 'Helvetica-Bold'  # Резервний жирний шрифт
    p.setFont(font_name, 12)
    
    # Налаштування стилю для Paragraph
    para_style = ParagraphStyle(
        name='Normal',
        fontName=font_name,
        fontSize=12,
        leading=14,  # Міжрядковий інтервал
        leftIndent=3 * cm,  # Відступ зліва
        spaceAfter=0.5 * cm,  # Відступ після абзацу
        allowWidows=1,
        allowOrphans=1,
    )
    
    # Початкові координати
    y = height - 2 * cm  # Відступ від верху сторінки
    
    # Додаємо фотографію (вгорі)
    if resume.photo:
        try:
            p.drawImage(resume.photo.path, 2 * cm, y - 5 * cm, width=5 * cm, height=5 * cm, preserveAspectRatio=True)
            y -= 6 * cm  # Зменшуємо y після фото
        except Exception as e:
            print(f"Error rendering image: {e}")
            p.drawString(2 * cm, y, "Не вдалося завантажити фотографію")
            y -= 1 * cm
    
    # Заголовок резюме
    p.setFont(font_bold_name, 16)
    p.drawString(2 * cm, y, resume.title[:100])  # Обмежуємо довжину заголовка
    y -= 1.5 * cm
    
    # Секції резюме
    for section in resume.sections.all().order_by('order'):
        # Заголовок секції (жирний шрифт)
        p.setFont(font_bold_name, 12)
        section_title = f"{section.get_section_type_display()}:"
        p.drawString(2 * cm, y, section_title)
        # Опція підкреслення 
        # text_width = p.stringWidth(section_title, font_bold_name, 12)
        # p.line(2 * cm, y - 0.2 * cm, 2 * cm + text_width, y - 0.2 * cm)
        y -= 0.3 * cm
        # Вміст секції
        content = section.content or ""
        if content:
            p.setFont(font_name, 12)
            # Створюємо Paragraph для автоматичного перенесення
            para = Paragraph(content.replace('\n', '<br/>'), para_style)
            # Обчислюємо необхідну висоту
            used_height = para.wrap(width - 4 * cm, y - 2 * cm)[1]
            if y - used_height < 2 * cm:  # Якщо не вистачає місця, нова сторінка
                p.showPage()
                p.setFont(font_bold_name, 12)
                y = height - 2 * cm
                p.drawString(2 * cm, y, section_title)
                # p.line(2 * cm, y - 0.2 * cm, 2 * cm + text_width, y - 0.2 * cm)  # Підкреслення на новій сторінці
                y -= 0.8 * cm
            # Малюємо Paragraph
            para.drawOn(p, 2 * cm, y - used_height)
            y -= used_height + 0.5 * cm  # Додаємо відступ після секції
            if y < 2 * cm:  # Нова сторінка, якщо закінчився простір
                p.showPage()
                p.setFont(font_name, 12)
                y = height - 2 * cm
        y -= 1 * cm  # Додатковий відступ
    
    p.showPage()
    p.save()
    return response


# В'ю для експорту конкретного резюме в docx-форматі
def export_docx(request, pk):
    resume = get_object_or_404(Resume, pk=pk, user=request.user)
    doc = Document()
    
    # Додаємо фотографію (вгорі)
    if resume.photo:
        try:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(resume.photo.path, width=Cm(5), height=Cm(5))
            paragraph.paragraph_format.space_after = Cm(0.5)
        except Exception as e:
            print(f"Error adding image to DOCX: {e}")
            doc.add_paragraph(f"Не вдалося завантажити фотографію")
    
    # Заголовок резюме
    doc.add_heading(resume.title[:100], 0)
    
    # Секції резюме
    for section in resume.sections.all().order_by('order'):
        # Налаштування стилю для жирного заголовка секції
        heading = doc.add_heading(level=1)
        run = heading.add_run(f"{section.get_section_type_display()}:")
        run.bold = True
        # Опція підкреслення
        # run.underline = True
        # Вміст секції
        doc.add_paragraph(section.content or "")
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{resume.title}.docx"'
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    response.write(doc_buffer.getvalue())
    return response


# --------- В'юшки для API ---------
# визначають, як дані будуть оброблятись і повертатись у відповідь на HTTP-запити

# Показ\додавання профілей
class ProfileListAPI(generics.ListCreateAPIView):
    queryset = Profile.objects.all()
    serializer_class = ProfileSerializer

# Редагування\видалення профілей
class ProfileDetailAPI(generics.RetrieveUpdateDestroyAPIView):
    queryset = Profile.objects.all()
    serializer_class = ProfileSerializer

# Показ\додавання шаблонів резюме
class ResumeTemplateListAPI(generics.ListCreateAPIView):
    queryset = ResumeTemplate.objects.all()
    serializer_class = ResumeTemplateSerializer

# Редагування\видалення шаблонів резюме
class ResumeTemplateDetailAPI(generics.RetrieveUpdateDestroyAPIView):
    queryset = ResumeTemplate.objects.all()
    serializer_class = ResumeTemplateSerializer

# Показ\додавання резюме
class ResumeListAPI(generics.ListCreateAPIView):
    queryset = Resume.objects.all()
    serializer_class = ResumeSerializer

# Редагування\видалення резюме
class ResumeDetailAPI(generics.RetrieveUpdateDestroyAPIView):
    queryset = Resume.objects.all()
    serializer_class = ResumeSerializer

# Показ\додавання секції резюме
class ResumeSectionListAPI(generics.ListCreateAPIView):
    queryset = ResumeSection.objects.all()
    serializer_class = ResumeSectionSerializer

# Редагування\видалення секції резюме
class ResumeSectionDetailAPI(generics.RetrieveUpdateDestroyAPIView):
    queryset = ResumeSection.objects.all()
    serializer_class = ResumeSectionSerializer

# Показ\додавання оголошень
class AnnouncementListAPI(generics.ListCreateAPIView):
    queryset = Announcement.objects.all()
    serializer_class = AnnouncementSerializer

# Редагування\видалення оголошень
class AnnouncementDetailAPI(generics.RetrieveUpdateDestroyAPIView):
    queryset = Announcement.objects.all()
    serializer_class = AnnouncementSerializer